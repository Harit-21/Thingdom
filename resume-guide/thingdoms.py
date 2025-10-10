# Ensure you have the library: pip install python-docx

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ==============================================================================
#  ✅ EDIT THIS SECTION TO CUSTOMIZE YOUR RESUME ✅
# ==============================================================================

# --- Fill in your personal and professional details below ---
resume_data = {
    # --- Personal Information ---
    "name": "First Last",
    "tagline": "Your Professional Title | e.g., Full-Stack Developer",
    "email": "your.email@example.com",
    "phone": "+1 (123) 456-7890",
    "location": "City, Country",
    "linkedin": "https://www.linkedin.com/in/yourprofile",
    "github": "https://github.com/yourusername",
    "portfolio": "https://yourportfolio.com",
    
    # --- Professional Summary ---
    # A brief, powerful overview of your skills and career goals.
    "summary": "Innovative Computer Science professional with hands-on experience in the full software development lifecycle (SDLC), from concept to deployment. A proactive problem-solver with a track record of building scalable, user-centric applications using agile methodologies. Passionate about leveraging modern technologies to solve complex challenges.",
    
    # --- Core Competencies ---
    # List your most valuable skills. Aim for 6 for a balanced look.
    "core_competencies": [
        "Full-Stack Development (MERN)", "Agile Project Management", "Data Structures & Algorithms",
        "UI/UX Design Principles", "Database Architecture", "Cloud Computing (AWS)",
        "CI/CD & DevOps", "API Design (RESTful)", "Mobile Development (React Native)"
    ],
    
    # --- Professional Experience ---
    # Add each job as a new dictionary item in the list.
    "experience": [
        {
            "title": "Software Engineer",
            "company": "Tech Solutions Inc.",
            "location": "San Francisco, CA",
            "date": "July 2025 – Present",
            "details": [
                'Lead the development of a key feature for a major product, resulting in a 15% increase in user engagement.',
                'Engineered a scalable backend service using Node.js, improving API response time by 40%.',
                'Collaborated in an agile team to manage the product lifecycle, from planning and sprints to deployment.',
                'Contributed to the development of a CI/CD pipeline, reducing deployment time by 50%.'
            ]
        },
        {
            "title": "Junior Developer",
            "company": "Creative Web Agency",
            "location": "Remote",
            "date": "Feb 2024 – July 2025",
            "details": [
                'Developed and launched 5+ custom websites for clients, increasing their online presence and sales.',
                'Implemented a content management system that saved clients over 10 hours of manual work per week.'
            ]
        }
    ],
    
    # --- Personal Projects ---
    # Showcase your personal work and passion projects.
    "projects": [
        {
            "name": "AI-Powered Task Manager",
            "type": "Personal Web Application",
            "location": "GitHub",
            "date": "Dec 2024",
            "details": [
                'Developed a React-based application that uses natural language processing to categorize and prioritize tasks.',
                'Implemented machine learning models to predict task completion times with 90% accuracy.'
            ]
        }
    ],
    
    # --- Education ---
    # List your degrees, starting with the most recent.
    "education": [
        {
            "degree": "Bachelor of Science in Computer Science",
            "institution": "State University",
            "location": "City, State",
            "date": "Expected Aug 2027",
            "details": [
                'Relevant Coursework: Advanced Algorithms, Data Structures, Software Engineering, Machine Learning.',
                'Academic Standing: Dean\'s List for four consecutive semesters.'
            ]
        }
    ]
}

# ==============================================================================
#  PRESENTATION LOGIC (NO NEED TO EDIT BELOW THIS LINE)
# ==============================================================================

# --- Helper Functions ---
def set_document_defaults(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x40, 0x40, 0x40)

def add_section_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.name = 'Garamond'
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    
    p_bdr = OxmlElement('w:pBdr')
    bottom_bdr = OxmlElement('w:bottom')
    bottom_bdr.set(qn('w:val'), 'single')
    bottom_bdr.set(qn('w:sz'), '4')
    bottom_bdr.set(qn('w:space'), '1')
    bottom_bdr.set(qn('w:color'), 'auto')
    p_bdr.append(bottom_bdr)
    p._p.get_or_add_pPr().append(p_bdr)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_entry(doc, primary_text, secondary_text, location, date_range, details):
    p1 = doc.add_paragraph()
    p1.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run_primary = p1.add_run(primary_text)
    run_primary.font.name = 'Garamond'
    run_primary.font.size = Pt(11)
    run_primary.bold = True
    p1.add_run(f'\t{date_range}')
    p1.paragraph_format.space_after = Pt(1)
    p2 = doc.add_paragraph()
    p2.add_run(secondary_text).italic = True
    p2.add_run(f' | {location}')
    p2.paragraph_format.space_after = Pt(2)
    for detail in details:
        p_bullet = doc.add_paragraph(style='List Bullet')
        p_bullet.add_run(detail)
        p_bullet.paragraph_format.left_indent = Inches(0.25)
        p_bullet.paragraph_format.space_after = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_core_competencies(doc, competencies):
    num_cols = 3
    num_rows = (len(competencies) + num_cols - 1) // num_cols
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    tbl_pr = table._tbl.tblPr
    tbl_look = OxmlElement('w:tblLook')
    tbl_look.set(qn('w:val'), '04A0')
    tbl_pr.append(tbl_look)
    
    competency_iterator = iter(competencies)
    for row in table.rows:
        for cell in row.cells:
            try:
                cell.text = next(competency_iterator)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except StopIteration:
                break
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# --- Main Generator Function ---
def generate_resume(data):
    """Creates a .docx resume from a dictionary of data."""
    doc = Document()
    set_document_defaults(doc)

    # Header
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(data["name"])
    name_run.font.name = 'Garamond'
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    tagline_p = doc.add_paragraph()
    tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline_p.add_run(data["tagline"]).bold = True
    tagline_p.paragraph_format.space_after = Pt(4)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.add_run('📧 ')
    add_hyperlink(contact_p, data["email"], f'mailto:{data["email"]}')
    contact_p.add_run(f'  |  📱 {data["phone"]}  |  ')
    add_hyperlink(contact_p, ' LinkedIn', data["linkedin"])
    contact_p.add_run('  |  ')
    add_hyperlink(contact_p, ' GitHub', data["github"])
    contact_p.add_run('  |  ')
    add_hyperlink(contact_p, '🌐 Portfolio', data["portfolio"])
    contact_p.paragraph_format.space_after = Pt(6)

    # Professional Summary
    summary_p = doc.add_paragraph()
    summary_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_p.add_run(data["summary"]).italic = True

    # Core Competencies
    add_section_header(doc, 'Core Competencies')
    add_core_competencies(doc, data["core_competencies"])

    # Professional Experience
    add_section_header(doc, 'Professional Experience')
    for job in data["experience"]:
        add_entry(doc, job["company"], job["title"], job["location"], job["date"], job["details"])

    # Personal Projects
    add_section_header(doc, 'Personal Projects')
    for project in data["projects"]:
        add_entry(doc, project["name"], project["type"], project["location"], project["date"], project["details"])
    
    # Education
    add_section_header(doc, 'Education')
    for edu in data["education"]:
        add_entry(doc, edu["institution"], edu["degree"], edu["location"], edu["date"], edu["details"])
    
    # Save the final document
    file_name = f"{data['name'].replace(' ', '_')}_Resume.docx"
    doc.save(file_name)
    print(f"✅ Successfully generated resume: {file_name}")

# --- Script Execution ---
if __name__ == "__main__":
    generate_resume(resume_data)